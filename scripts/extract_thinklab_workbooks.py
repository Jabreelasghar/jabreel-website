from pathlib import Path
from docx import Document


SOURCES = {
    "session-1": Path(
        r"C:\Users\jabre\OneDrive - Higher Colleges of Technology\PERSONAL\ThinkLab"
        r"\ThinkLab™ Guided Experience 1 — The Trust Lab"
        r"\ThinkLab_Guided_Experience_Student_Edition.docx"
    ),
    "session-2": Path(
        r"C:\Users\jabre\OneDrive - Higher Colleges of Technology\PERSONAL\ThinkLab"
        r"\ThinkLab™ Guided Experience 2 — The Decision Lab"
        r"\SESSION 2 – THE DECISION LAB_2.docx"
    ),
}


def cell_text(cell) -> str:
    return " / ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def extract(name: str, source: Path) -> None:
    doc = Document(source)
    lines = [f"# {name}", "", f"Source: {source}", ""]
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(f"[{paragraph.style.name}] {text}")
    for index, table in enumerate(doc.tables, start=1):
        lines.extend(["", f"## Table {index}", ""])
        for row in table.rows:
            lines.append(" | ".join(cell_text(cell) for cell in row.cells))
    output = Path("work/thinklab") / f"{name}-structure.md"
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text("\n".join(lines), encoding="utf-8")


for workbook_name, workbook_path in SOURCES.items():
    extract(workbook_name, workbook_path)
